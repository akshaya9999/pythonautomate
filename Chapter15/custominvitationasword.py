import docx

doc=docx.Document()
file=open('guests.txt')
n=[]
n=file.readlines()
for i in n:
    a=doc.add_paragraph("")
    a.alignment=1
    a.add_run('It would be the pleasure to have the company of').italic=True
    x=doc.add_paragraph("")
    x.alignment=1
    x.add_run(i[:-1]).bold=True
    b=doc.add_paragraph("")
    b.alignment=1
    b.add_run('at 11011 Memory Lane on the Evening of').italic=True
    c=doc.add_paragraph("")
    c.alignment=1
    c.add_run('April 1st').bold=True
    d=doc.add_paragraph("")
    d.alignment=1
    d.add_run('at 7 o\' clock').italic=True
    doc.add_page_break()
doc.save('custominv.docx')

